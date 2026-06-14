#!/usr/bin/env python3
"""python-docx report skeleton — copy and customize for each report.

Pitfalls documented in the skill: use system Python (not Hermes venv), 
write to file first (not inline in terminal), use Light Grid Accent 1 for tables.
"""

import os, datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.size = Pt(11)

# === Cover ===
for _ in range(3):
    doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Report Title\nSubtitle')
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run(f'Version: v1.0 | Date: {datetime.date.today().isoformat()}')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
doc.add_page_break()

# === TOC ===
doc.add_heading('目录', level=1)
for item in ['1. Background', '2. Analysis', '3. Conclusion', 'A. Appendix']:
    doc.add_paragraph(item)

# === Table helper ===
def make_table(doc, headers, data):
    t = doc.add_table(rows=len(data)+1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for run in t.rows[0].cells[i].paragraphs[0].runs:
            run.font.bold = True
    for ri, rd in enumerate(data, 1):
        for ci, v in enumerate(rd):
            t.rows[ri].cells[ci].text = v
    return t

# === Code block helper ===
def code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return p

# ====== YOUR REPORT CONTENT HERE ======
doc.add_heading('1. Background', level=1)
doc.add_paragraph('Replace with actual content.')

# Example table
make_table(doc, ['Col A', 'Col B', 'Col C'], [
    ['val1', 'val2', 'val3'],
])

# ====== Save ======
output = os.path.expanduser('~/Desktop/report.docx')
doc.save(output)
print(f'✅ Saved: {output}')
